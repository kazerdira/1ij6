#!/usr/bin/env python3
"""
Export Utilities for Real-time Speech Translator
Supports SRT, VTT, JSON, TXT, and other formats
"""

import json
from datetime import timedelta
from typing import List, Dict
import xml.etree.ElementTree as ET


class ExportManager:
    """Manage export of transcriptions to various formats"""
    
    @staticmethod
    def format_timestamp_srt(seconds: float) -> str:
        """
        Format seconds to SRT timestamp (HH:MM:SS,mmm)
        
        Args:
            seconds: Time in seconds
        
        Returns:
            SRT formatted timestamp
        """
        td = timedelta(seconds=seconds)
        hours = int(td.total_seconds() // 3600)
        minutes = int((td.total_seconds() % 3600) // 60)
        secs = int(td.total_seconds() % 60)
        millis = int((td.total_seconds() % 1) * 1000)
        
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"
    
    @staticmethod
    def format_timestamp_vtt(seconds: float) -> str:
        """
        Format seconds to WebVTT timestamp (HH:MM:SS.mmm)
        
        Args:
            seconds: Time in seconds
        
        Returns:
            VTT formatted timestamp
        """
        td = timedelta(seconds=seconds)
        hours = int(td.total_seconds() // 3600)
        minutes = int((td.total_seconds() % 3600) // 60)
        secs = int(td.total_seconds() % 60)
        millis = int((td.total_seconds() % 1) * 1000)
        
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"
    
    def export_srt(self, transcriptions: List[Dict], output_path: str):
        """
        Export transcriptions to SRT subtitle format
        
        Args:
            transcriptions: List of transcription dictionaries
            output_path: Output file path
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            for i, trans in enumerate(transcriptions, 1):
                start = self.format_timestamp_srt(trans['start'])
                end = self.format_timestamp_srt(trans['end'])
                
                f.write(f"{i}\n")
                f.write(f"{start} --> {end}\n")
                f.write(f"{trans.get('original', '')}\n")
                f.write(f"{trans.get('translated', '')}\n")
                f.write("\n")
    
    def export_vtt(self, transcriptions: List[Dict], output_path: str):
        """
        Export transcriptions to WebVTT subtitle format
        
        Args:
            transcriptions: List of transcription dictionaries
            output_path: Output file path
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("WEBVTT\n\n")
            
            for i, trans in enumerate(transcriptions, 1):
                start = self.format_timestamp_vtt(trans['start'])
                end = self.format_timestamp_vtt(trans['end'])
                
                f.write(f"{i}\n")
                f.write(f"{start} --> {end}\n")
                f.write(f"{trans.get('original', '')}\n")
                f.write(f"{trans.get('translated', '')}\n")
                f.write("\n")
    
    def export_json(self, transcriptions: List[Dict], output_path: str, metadata: Dict = None):
        """
        Export transcriptions to JSON format
        
        Args:
            transcriptions: List of transcription dictionaries
            output_path: Output file path
            metadata: Additional metadata to include
        """
        data = {
            "transcriptions": transcriptions
        }
        
        if metadata:
            data["metadata"] = metadata
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    
    def export_txt(self, transcriptions: List[Dict], output_path: str, include_timestamps: bool = True):
        """
        Export transcriptions to plain text format
        
        Args:
            transcriptions: List of transcription dictionaries
            output_path: Output file path
            include_timestamps: Include timestamps in output
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            for trans in transcriptions:
                if include_timestamps:
                    f.write(f"[{trans['start']:.2f}s - {trans['end']:.2f}s]\n")
                
                f.write(f"Original: {trans.get('original', '')}\n")
                f.write(f"Translated: {trans.get('translated', '')}\n")
                f.write("\n")
    
    def export_xml(self, transcriptions: List[Dict], output_path: str):
        """
        Export transcriptions to XML format
        
        Args:
            transcriptions: List of transcription dictionaries
            output_path: Output file path
        """
        root = ET.Element("transcriptions")
        
        for trans in transcriptions:
            segment = ET.SubElement(root, "segment")
            
            start = ET.SubElement(segment, "start")
            start.text = str(trans['start'])
            
            end = ET.SubElement(segment, "end")
            end.text = str(trans['end'])
            
            original = ET.SubElement(segment, "original")
            original.text = trans.get('original', '')
            
            translated = ET.SubElement(segment, "translated")
            translated.text = trans.get('translated', '')
        
        tree = ET.ElementTree(root)
        tree.write(output_path, encoding='utf-8', xml_declaration=True)
    
    def export_csv(self, transcriptions: List[Dict], output_path: str):
        """
        Export transcriptions to CSV format
        
        Args:
            transcriptions: List of transcription dictionaries
            output_path: Output file path
        """
        import csv
        
        with open(output_path, 'w', encoding='utf-8', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(['Start', 'End', 'Original', 'Translated'])
            
            for trans in transcriptions:
                writer.writerow([
                    trans['start'],
                    trans['end'],
                    trans.get('original', ''),
                    trans.get('translated', '')
                ])
    
    def export_markdown(self, transcriptions: List[Dict], output_path: str, title: str = "Transcription"):
        """
        Export transcriptions to Markdown format
        
        Args:
            transcriptions: List of transcription dictionaries
            output_path: Output file path
            title: Document title
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"# {title}\n\n")
            f.write("---\n\n")
            
            for i, trans in enumerate(transcriptions, 1):
                f.write(f"## Segment {i}\n\n")
                f.write(f"**Time:** {trans['start']:.2f}s - {trans['end']:.2f}s\n\n")
                f.write(f"**Original:**\n> {trans.get('original', '')}\n\n")
                f.write(f"**Translated:**\n> {trans.get('translated', '')}\n\n")
                f.write("---\n\n")
    
    def export_docx(self, transcriptions: List[Dict], output_path: str, title: str = "Transcription"):
        """
        Export transcriptions to DOCX format
        
        Args:
            transcriptions: List of transcription dictionaries
            output_path: Output file path
            title: Document title
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError:
            raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")
        
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        
        # Add transcriptions
        for i, trans in enumerate(transcriptions, 1):
            # Segment header
            doc.add_heading(f"Segment {i}", level=2)
            
            # Timestamp
            time_para = doc.add_paragraph()
            time_run = time_para.add_run(f"Time: {trans['start']:.2f}s - {trans['end']:.2f}s")
            time_run.font.size = Pt(10)
            time_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Original
            orig_para = doc.add_paragraph()
            orig_para.add_run("Original: ").bold = True
            orig_para.add_run(trans.get('original', ''))
            
            # Translated
            trans_para = doc.add_paragraph()
            trans_para.add_run("Translated: ").bold = True
            trans_para.add_run(trans.get('translated', ''))
            
            # Separator
            doc.add_paragraph()
        
        doc.save(output_path)
    
    def export_all(self, transcriptions: List[Dict], base_path: str, metadata: Dict = None):
        """
        Export transcriptions to all supported formats
        
        Args:
            transcriptions: List of transcription dictionaries
            base_path: Base path for output files (without extension)
            metadata: Additional metadata
        """
        formats = {
            'srt': self.export_srt,
            'vtt': self.export_vtt,
            'json': lambda t, p: self.export_json(t, p, metadata),
            'txt': self.export_txt,
            'xml': self.export_xml,
            'csv': self.export_csv,
            'md': self.export_markdown
        }
        
        for ext, export_func in formats.items():
            try:
                output_path = f"{base_path}.{ext}"
                export_func(transcriptions, output_path)
                print(f"✅ Exported: {output_path}")
            except Exception as e:
                print(f"❌ Failed to export {ext}: {e}")


def main():
    """Test export utilities"""
    # Sample data
    transcriptions = [
        {
            "start": 0.0,
            "end": 3.5,
            "original": "안녕하세요",
            "translated": "Hello"
        },
        {
            "start": 3.5,
            "end": 7.2,
            "original": "오늘 날씨가 좋네요",
            "translated": "The weather is nice today"
        },
        {
            "start": 7.2,
            "end": 11.8,
            "original": "만나서 반갑습니다",
            "translated": "Nice to meet you"
        }
    ]
    
    exporter = ExportManager()
    
    print("Testing export utilities...")
    print()
    
    # Test all formats
    metadata = {
        "source_language": "ko",
        "target_language": "eng_Latn",
        "model": "base"
    }
    
    exporter.export_all(transcriptions, "test_output", metadata)
    
    print()
    print("Export test complete!")


if __name__ == "__main__":
    main()
